from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs'

BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
INK = '1F2937'
MUTED = '5B6470'
LIGHT = 'F2F4F7'
GREEN = 'EAF4EE'
AMBER = 'FFF5E5'


def set_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = 'STHeiti'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'STHeiti')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'STHeiti')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'STHeiti')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def table(doc, headers, rows, widths, status_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    t.style = 'Table Grid'
    for idx, (cell, text, width) in enumerate(zip(t.rows[0].cells, headers, widths)):
        cell.width = Inches(width)
        shade(cell, LIGHT)
        set_cell_margin(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, 9, DARK_BLUE, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for idx, (cell, value, width) in enumerate(zip(cells, row, widths)):
            cell.width = Inches(width)
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if status_col == idx:
                if str(value) == 'passed': shade(cell, GREEN)
                elif str(value) in ('manual_required', 'environment_blocked'): shade(cell, AMBER)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_font(r, 8.5, INK, bold=(status_col == idx))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def setup(doc, title, subtitle):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    normal = doc.styles['Normal']
    normal.font.name = 'STHeiti'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'STHeiti')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'STHeiti')
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'STHeiti')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in [('Heading 1', 16, BLUE, 16, 8), ('Heading 2', 13, BLUE, 12, 6), ('Heading 3', 12, DARK_BLUE, 8, 4)]:
        style = doc.styles[style_name]
        style.font.name = 'STHeiti'; style._element.rPr.rFonts.set(qn('w:eastAsia'), 'STHeiti'); style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run('AgentLab v0.3.0 · Quality Assurance')
    set_font(r, 8, MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run('Internal test artifact · 2026-08-01')
    set_font(r, 8, MUTED)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title); set_font(r, 23, '000000', bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle); set_font(r, 12, MUTED)
    meta = [('测试对象', 'CodexMVP / AgentLab v0.3.0'), ('代码基线', '1193dce · tag v0.3.0'), ('执行日期', '2026-08-01')]
    for label, value in meta:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f'{label}: '); set_font(r, 10, INK, bold=True)
        r = p.add_run(value); set_font(r, 10, INK)


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead); set_font(r, 11, INK, bold=True)
        r = p.add_run(text[len(bold_lead):]); set_font(r)
    else:
        r = p.add_run(text); set_font(r)
    return p


def build_plan():
    doc = Document()
    setup(doc, 'v0.3.0 测试方案', '安全编码交付闭环 · 自动化、候选包与人工验收')
    doc.add_heading('1. 目标与范围', level=1)
    add_para(doc, '验证 v0.3.0 是否形成“安全执行 → 变更审查 → 验证 → 可追溯提交”的最小闭环，并优先回归 V2 报告确认的 TC-10 生命周期和只读权限风险。')
    add_para(doc, '离线范围包括 SQLite 状态恢复、Git 服务、验证服务、Sandbox 预设、权限审计、多 Agent 策略、构建和候选包。真实 LLM、MCP/ROS、人工 GUI 与签名公证不以自动化通过替代。')
    doc.add_heading('2. 准入与判定', level=1)
    table(doc, ['项目', '准入/判定'], [
        ('代码基线', '1193dce07b67668983df34bc7be20c2d3504b47c，v0.3.0'),
        ('环境', 'macOS arm64；Node.js 24.x；Electron 43.1.1；本地 Git 可用'),
        ('自动化通过', '命令退出码为 0 或全部断言通过'),
        ('阻塞分类', '缺失外部凭据、服务、设备或人工 GUI 前置条件，标记 environment_blocked/manual_required'),
    ], [1.55, 4.95])
    doc.add_heading('3. 用例矩阵', level=1)
    plan_rows = [
        ('TC-01', 'P0', '全量回归', 'npm test', '所有测试通过'),
        ('TC-02', 'P0', 'TypeScript/前端/主进程构建', 'npm run build:web', '构建成功'),
        ('TC-03', 'P0', 'ESLint 9 Flat Config', 'npm run lint', '零 error/zero warning'),
        ('TC-04', 'P0', 'SQLite 恢复', 'SessionStore 测试', 'running Turn 恢复 interrupted'),
        ('TC-05', 'P0', 'Git 乐观并发', 'GitService 测试', '旧 token 被拒绝'),
        ('TC-06', 'P0', '只读权限继承', 'policy 测试', '危险工具全部拒绝'),
        ('TC-07', 'P0', 'TC-10 终态规则', 'policy ×10', '只认可四种终态'),
        ('TC-08', 'P0', '事件/权限/SDK UUID', '编译与代码核验', '持久化路径完整'),
        ('TC-09', 'P0', '候选应用结构', 'Info.plist/二进制', '版本 0.3.0'),
        ('TC-10~13', 'P1', 'Git UI、Review、Sandbox、真实三专家', '人工 GUI / 有效 LLM', '按 PRD 验收'),
        ('TC-14', 'P2', 'MCP/ROS/车辆扩展', '专用环境', '缺环境不误报通过'),
    ]
    table(doc, ['ID', '级别', '场景', '方法', '通过标准'], plan_rows, [0.55, 0.45, 1.55, 1.65, 2.3])
    doc.add_heading('4. 执行顺序', level=1)
    for item in ['执行 npm test，读取测试数量与失败数。', '执行 npm run lint 与 npm run build:web。', '核验候选 .app 的 Info.plist、主可执行文件和版本号。', '将真实 LLM/GUI/MCP 场景交给有凭据和隔离工作区的人工补测。']:
        p = doc.add_paragraph(style='List Number'); r = p.add_run(item); set_font(r)
    doc.add_heading('5. 发布门', level=1)
    add_para(doc, '所有 P0 自动化项必须通过。未签名候选应用只能用于本地验收；正式 macOS 分发前必须完成 Developer ID 签名、notarization 与 Gatekeeper 验证。')
    doc.save(OUT / 'CodexMVP-v0.3.0-测试方案.docx')


def build_report():
    doc = Document()
    setup(doc, 'v0.3.0 测试结果报告', '执行结论 · 自动化发布门通过，待进入人工验收')
    doc.add_heading('1. 执行结论', level=1)
    add_para(doc, '自动化发布门通过，候选包可进入本地人工验收；不建议将未签名候选包作为对外 macOS 发布物。')
    add_para(doc, '本次执行取得：16/16 自动化测试通过、ESLint 通过、TypeScript/Web/Electron 构建通过，候选应用版本为 0.3.0。V2 阻塞项中的 TC-10 终态规则、只读越权、SQLite 恢复与 ESLint 9 均有实现和自动化证据。')
    doc.add_heading('2. 环境与命令证据', level=1)
    table(doc, ['项目', '实际结果'], [
        ('代码', '1193dce / v0.3.0'), ('平台', 'macOS arm64'), ('测试命令', 'npm test'),
        ('静态检查', 'npm run lint'), ('构建', 'npm run build:web'), ('候选包', 'release-candidate/mac-arm64/AgentLab.app'),
    ], [1.55, 4.95])
    doc.add_heading('3. 结果汇总', level=1)
    table(doc, ['状态', '数量', '解释'], [('passed', '9', 'P0 自动化与候选包核验通过'), ('manual_required', '3', '需人工 GUI 或真实 LLM 凭据'), ('environment_blocked', '2', '真实 LLM 与 MCP/ROS/车辆环境未配置'), ('failed', '0', '无失败项')], [1.45, 0.7, 4.35], status_col=0)
    doc.add_heading('4. 用例结果', level=1)
    result_rows = [
        ('TC-01', 'passed', 'npm test：16 tests / 9 suites / 0 fail'),
        ('TC-02', 'passed', 'npm run build:web 成功；TypeScript 无错误'),
        ('TC-03', 'passed', 'npm run lint 成功；零 warning'),
        ('TC-04', 'passed', 'SessionStore 2 项通过；重启后 running Turn 变 interrupted'),
        ('TC-05', 'passed', 'GitService 暂存成功；过期 state token 被拒绝'),
        ('TC-06', 'passed', '只读 Reviewer 拒绝 Bash/Edit/Write/Web/未声明 MCP'),
        ('TC-07', 'passed', '10 轮策略断言通过；仅四种 SDK 终态可完成'),
        ('TC-08', 'passed', 'Turn/Event/Permission/Subagent/Verification SQLite 路径编译通过'),
        ('TC-09', 'passed', 'Info.plist 与主可执行文件存在；CFBundleShortVersionString=0.3.0'),
        ('TC-10', 'manual_required', '待人工核对 Git UI 与命令行一致性、Revert 二次确认'),
        ('TC-11', 'manual_required', '待真实 LLM 验证 Review、验证失败转 Prompt'),
        ('TC-12', 'manual_required', '待真实工具调用验证 Sandbox 与权限弹窗'),
        ('TC-13', 'environment_blocked', '本次无可用 LLM/API 凭据，未运行真实三专家连续 10 次'),
        ('TC-14', 'environment_blocked', '本机未提供 MCP/ROS/车辆依赖'),
    ]
    table(doc, ['ID', '结果', '证据/说明'], result_rows, [0.65, 1.35, 4.5], status_col=1)
    doc.add_heading('5. V2 风险回归', level=1)
    table(doc, ['V2 风险', 'v0.3.0 处理与证据'], [
        ('主任务未等待专家终态', '持久化 SDK task_* 与 background_tasks_changed；未终态任务令 Turn 失败/不完整；终态策略测试通过。'),
        ('只读专家越权', 'agentID 映射工具白名单，未知身份降级只读；10 轮拒绝危险工具测试通过。'),
        ('状态恢复不足', 'SQLite 保存 Thread/Turn/Event/Permission/Verification/Subagent；恢复测试通过。'),
        ('Lint 失败与资产缺失', 'ESLint 9 Flat Config 已通过；TC-10 固定 fixture 已纳入仓库。'),
    ], [1.65, 4.85])
    doc.add_heading('6. 风险与后续动作', level=1)
    for item in ['完成 Developer ID 签名、notarization 和 Gatekeeper 验证后再进行外部分发。', '配置有效 LLM API，在隔离 Git 工作区执行 TC-10~TC-13；真实三专家场景连续执行 10 次。', '人工验收 Hunk Git 操作、Sandbox 越界阻止与权限弹窗信息准确性。', '在专用 Fixture 环境补测 MCP/ROS/车辆扩展用例，继续区分环境阻塞与功能失败。']:
        p = doc.add_paragraph(style='List Number'); r = p.add_run(item); set_font(r)
    doc.save(OUT / 'CodexMVP-v0.3.0-测试结果报告.docx')


if __name__ == '__main__':
    OUT.mkdir(parents=True, exist_ok=True)
    build_plan()
    build_report()
